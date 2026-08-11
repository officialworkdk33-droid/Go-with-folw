"""
Core logic for the HAWB Document Merger:
 - extract text from mixed file types
 - detect a HAWB (House Airway Bill) number in that text
 - merge a group of files into one PDF
"""
import io
import os
import re
import traceback
from pathlib import Path

SUPPORTED_EXTS = {
    "pdf", "docx", "doc", "xlsx", "xls", "csv", "txt",
    "png", "jpg", "jpeg", "webp", "bmp", "gif", "tif", "tiff",
}
IMAGE_EXTS = {"png", "jpg", "jpeg", "webp", "bmp", "gif", "tif", "tiff"}


def get_ext(filename: str) -> str:
    m = re.search(r"\.([A-Za-z0-9]+)$", filename)
    return m.group(1).lower() if m else ""


def sanitize_name(s: str) -> str:
    s = (s or "").strip()
    s = re.sub(r'[\\/:*?"<>|]+', "-", s)
    s = re.sub(r"\s+", "-", s)
    s = re.sub(r"-+", "-", s)
    return s or "UNNAMED"


def sanitize_key(k: str) -> str:
    # Strip ALL punctuation (not just disallowed chars) so the same HAWB written
    # as "88213/AX", "88213-AX" or "88213 AX" in different documents still
    # collapses to one canonical grouping key.
    k = re.sub(r"[^A-Za-z0-9]", "", k or "")
    return k.upper()


# ---------------------------------------------------------------------------
# Text extraction
# ---------------------------------------------------------------------------

def extract_text(path: Path, ext: str, log=None) -> str:
    try:
        if ext == "pdf":
            return _extract_pdf(path)
        if ext == "docx":
            return _extract_docx(path)
        if ext == "doc":
            return _extract_doc_legacy(path)
        if ext in ("xlsx", "xls"):
            return _extract_excel(path, ext)
        if ext in ("csv", "txt"):
            return path.read_text(errors="ignore")
        if ext in IMAGE_EXTS:
            if log:
                log(f"OCR reading {path.name} …")
            return _extract_image_ocr(path)
        return ""
    except Exception as e:
        if log:
            log(f"Could not read {path.name}: {e}", err=True)
        return ""


def _extract_pdf(path: Path) -> str:
    text_parts = []
    try:
        import pdfplumber
        with pdfplumber.open(str(path)) as pdf:
            for page in pdf.pages:
                t = page.extract_text() or ""
                text_parts.append(t)
        joined = "\n".join(text_parts)
        if joined.strip():
            return joined
    except Exception:
        pass
    # fallback to pypdf if pdfplumber found nothing / failed
    try:
        from pypdf import PdfReader
        reader = PdfReader(str(path))
        return "\n".join((p.extract_text() or "") for p in reader.pages)
    except Exception:
        return ""


def _extract_docx(path: Path) -> str:
    import docx
    d = docx.Document(str(path))
    parts = [p.text for p in d.paragraphs]
    for table in d.tables:
        for row in table.rows:
            parts.append(" | ".join(c.text for c in row.cells))
    return "\n".join(parts)


def _extract_doc_legacy(path: Path) -> str:
    # Old binary .doc: best-effort raw scan for readable ASCII runs.
    raw = path.read_bytes().decode("latin1", errors="ignore")
    cleaned = re.sub(r"[^\x20-\x7E\n]+", " ", raw)
    return cleaned


def _extract_excel(path: Path, ext: str) -> str:
    out = []
    if ext == "xlsx":
        import openpyxl
        wb = openpyxl.load_workbook(str(path), data_only=True)
        for ws in wb.worksheets:
            for row in ws.iter_rows(values_only=True):
                out.append(" | ".join("" if c is None else str(c) for c in row))
    else:  # legacy .xls
        try:
            import xlrd
            wb = xlrd.open_workbook(str(path))
            for sheet in wb.sheets():
                for r in range(sheet.nrows):
                    out.append(" | ".join(str(c) for c in sheet.row_values(r)))
        except Exception:
            return ""
    return "\n".join(out)


def _extract_image_ocr(path: Path) -> str:
    import pytesseract
    from PIL import Image
    with Image.open(path) as img:
        return pytesseract.image_to_string(img)


# ---------------------------------------------------------------------------
# HAWB detection
# ---------------------------------------------------------------------------

## The number-token group requires a digit somewhere in it (via lookahead) so that
## plain words like "House" or "Bill" that appear between the label and the actual
## number can never be mistaken for the HAWB number itself.
_TOKEN = r"(?=[A-Z0-9\-/]*\d)[A-Z0-9][A-Z0-9\-/]{3,19}"

_LABEL_RE = re.compile(
    r"(HOUSE\s*(?:AIR\s*WAY\s*BILL|AIRWAY\s*BILL|WAY\s*BILL)|H\.?A\.?W\.?B\.?|HAWB|HBL)"
    r".{0,40}?(" + _TOKEN + r")",
    re.IGNORECASE | re.DOTALL,
)
_AWB_RE = re.compile(r"\bAWB.{0,20}?(" + _TOKEN + r")", re.IGNORECASE | re.DOTALL)
_MAWB_RE = re.compile(r"\b(\d{3}-\d{7,8})\b")
_FILENAME_RE = re.compile(r"([A-Z0-9]{2,4}-?\d{6,12})", re.IGNORECASE)


def detect_hawb(text: str, filename: str):
    candidates = []
    for m in _LABEL_RE.finditer(text or ""):
        candidates.append((m.group(2), 1))
    if not candidates:
        for m in _AWB_RE.finditer(text or ""):
            candidates.append((m.group(1), 2))
    if not candidates:
        for m in _MAWB_RE.finditer(text or ""):
            candidates.append((m.group(1), 3))
    if not candidates:
        m = _FILENAME_RE.search(filename or "")
        if m:
            candidates.append((m.group(1), 4))
    if not candidates:
        return None
    candidates.sort(key=lambda c: c[1])
    return sanitize_key(candidates[0][0])


# ---------------------------------------------------------------------------
# PDF merge
# ---------------------------------------------------------------------------

def _text_to_pdf_bytes(title: str, text: str) -> bytes:
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import LETTER
    from reportlab.pdfbase.pdfmetrics import stringWidth

    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=LETTER)
    page_w, page_h = LETTER
    margin = 48
    size = 9.5
    line_height = 13
    max_width = page_w - margin * 2

    def safe(t):
        return (t or "").encode("latin-1", errors="replace").decode("latin-1")

    c.setFont("Helvetica-Bold", 13)
    c.drawString(margin, page_h - margin, safe(title)[:90])
    c.setFont("Helvetica", size)
    y = page_h - margin - 24

    for para in safe(text).split("\n"):
        words = para.split()
        line = ""
        lines = []
        if not words:
            lines.append("")
        for w in words:
            test = f"{line} {w}".strip()
            if stringWidth(test, "Helvetica", size) > max_width and line:
                lines.append(line)
                line = w
            else:
                line = test
        if line:
            lines.append(line)
        for ln in lines:
            if y < margin:
                c.showPage()
                c.setFont("Helvetica", size)
                y = page_h - margin
            c.drawString(margin, y, ln)
            y -= line_height
    c.showPage()
    c.save()
    return buf.getvalue()


def _image_to_pdf_bytes(path: Path) -> bytes:
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import LETTER
    from PIL import Image

    buf = io.BytesIO()
    page_w, page_h = LETTER
    margin = 30
    with Image.open(path) as img:
        img = img.convert("RGB")
        iw, ih = img.size
        scale = min((page_w - margin * 2) / iw, (page_h - margin * 2) / ih, 1.0)
        w, h = iw * scale, ih * scale
        c = canvas.Canvas(buf, pagesize=LETTER)
        c.drawImage(
            ImageReaderCompat(img), (page_w - w) / 2, (page_h - h) / 2, width=w, height=h
        )
        c.showPage()
        c.save()
    return buf.getvalue()


def ImageReaderCompat(pil_img):
    from reportlab.lib.utils import ImageReader
    return ImageReader(pil_img)


def build_merged_pdf(files, out_path: Path, log=None):
    """
    files: list of dicts {path: Path, name: str, ext: str, text: str}
    Writes a merged PDF to out_path.
    """
    from pypdf import PdfReader, PdfWriter

    writer = PdfWriter()

    for f in files:
        ext = f["ext"]
        try:
            if ext == "pdf":
                reader = PdfReader(str(f["path"]))
                for page in reader.pages:
                    writer.add_page(page)
            elif ext in IMAGE_EXTS:
                pdf_bytes = _image_to_pdf_bytes(f["path"])
                reader = PdfReader(io.BytesIO(pdf_bytes))
                for page in reader.pages:
                    writer.add_page(page)
            else:
                text = (f.get("text") or "").strip()
                note = text if text else "[No extractable text for this file type — the original is kept in this subfolder.]"
                pdf_bytes = _text_to_pdf_bytes(f["name"], note)
                reader = PdfReader(io.BytesIO(pdf_bytes))
                for page in reader.pages:
                    writer.add_page(page)
            if log:
                log(f"{f['name']} → merged")
        except Exception as e:
            if log:
                log(f"{f['name']} → merge failed, added as note page ({e})", err=True)
            try:
                pdf_bytes = _text_to_pdf_bytes(
                    f["name"],
                    f"[Could not merge this file: {e}\n{traceback.format_exc(limit=2)}\n"
                    "The original is kept in this subfolder.]",
                )
                reader = PdfReader(io.BytesIO(pdf_bytes))
                for page in reader.pages:
                    writer.add_page(page)
            except Exception:
                pass

    if len(writer.pages) == 0:
        # Ensure a valid (non-empty) PDF is always produced
        pdf_bytes = _text_to_pdf_bytes("Empty group", "No files could be merged into this PDF.")
        reader = PdfReader(io.BytesIO(pdf_bytes))
        for page in reader.pages:
            writer.add_page(page)

    with open(out_path, "wb") as fh:
        writer.write(fh)
